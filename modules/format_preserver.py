"""
日语到中文翻译程序 - 格式保存模块

此模块负责保存翻译后文档的格式，使其与原文档格式一致
"""

import os
import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import PyPDF2
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from typing import List, Dict, Any, Tuple

class FormatPreserver:
    """格式保存器，负责保持翻译后文档的格式与原文档一致"""
    
    def __init__(self, preserve_format: bool = True):
        """
        初始化格式保存器
        
        Args:
            preserve_format: 是否保存格式
        """
        self.preserve_format = preserve_format
    
    def create_document(self, original_paragraphs: List[str], translated_paragraphs: List[str], 
                       format_info: Dict[str, Any], output_path: str, file_type: str) -> bool:
        """
        创建格式一致的文档
        
        Args:
            original_paragraphs: 原文段落列表
            translated_paragraphs: 翻译后段落列表
            format_info: 格式信息
            output_path: 输出文件路径
            file_type: 文件类型，'docx'或'pdf'
            
        Returns:
            bool: 是否创建成功
        """
        if not self.preserve_format:
            # 如果不需要保存格式，直接创建纯文本文档
            return self._create_plain_document(translated_paragraphs, output_path, file_type)
        
        if file_type == 'docx':
            return self._create_docx_document(original_paragraphs, translated_paragraphs, format_info, output_path)
        elif file_type == 'pdf':
            return self._create_pdf_document(original_paragraphs, translated_paragraphs, format_info, output_path)
        else:
            raise ValueError(f"不支持的文件类型: {file_type}")
    
    def _create_plain_document(self, translated_paragraphs: List[str], output_path: str, file_type: str) -> bool:
        """
        创建纯文本文档
        
        Args:
            translated_paragraphs: 翻译后段落列表
            output_path: 输出文件路径
            file_type: 文件类型
            
        Returns:
            bool: 是否创建成功
        """
        try:
            if file_type == 'docx':
                # 创建Word文档
                doc = docx.Document()
                for para in translated_paragraphs:
                    doc.add_paragraph(para)
                doc.save(output_path)
            elif file_type == 'pdf':
                # 创建PDF文档
                c = canvas.Canvas(output_path, pagesize=letter)
                width, height = letter
                y = height - 50  # 起始y坐标
                
                # 注册中文字体
                try:
                    pdfmetrics.registerFont(TTFont('SimSun', 'simsun.ttf'))
                    font_name = 'SimSun'
                except:
                    font_name = 'Helvetica'  # 如果没有中文字体，使用默认字体
                
                for para in translated_paragraphs:
                    # 简单的文本换行处理
                    words = para.split()
                    line = ""
                    for word in words:
                        if len(line + word) > 80:  # 简单的行宽限制
                            c.setFont(font_name, 12)
                            c.drawString(50, y, line)
                            line = word + " "
                            y -= 15
                            if y < 50:  # 如果到达页面底部，创建新页面
                                c.showPage()
                                y = height - 50
                        else:
                            line += word + " "
                    
                    if line:
                        c.setFont(font_name, 12)
                        c.drawString(50, y, line)
                        y -= 30  # 段落间距
                    
                    if y < 50:  # 如果到达页面底部，创建新页面
                        c.showPage()
                        y = height - 50
                
                c.save()
            else:
                # 创建纯文本文件
                with open(output_path, 'w', encoding='utf-8') as f:
                    for para in translated_paragraphs:
                        f.write(para + '\n\n')
            
            return True
        except Exception as e:
            print(f"创建纯文本文档失败: {e}")
            return False
    
    def _create_docx_document(self, original_paragraphs: List[str], translated_paragraphs: List[str], 
                             format_info: Dict[str, Any], output_path: str) -> bool:
        """
        创建格式一致的Word文档
        
        Args:
            original_paragraphs: 原文段落列表
            translated_paragraphs: 翻译后段落列表
            format_info: 格式信息
            output_path: 输出文件路径
            
        Returns:
            bool: 是否创建成功
        """
        try:
            # 创建新文档
            doc = docx.Document()
            
            # 应用格式信息
            styles = format_info.get('styles', [])
            
            # 确保翻译后的段落数量与原文一致
            if len(original_paragraphs) != len(translated_paragraphs):
                print(f"警告: 翻译后段落数量({len(translated_paragraphs)})与原文({len(original_paragraphs)})不一致")
            
            # 添加段落并应用格式
            for i, translated_para in enumerate(translated_paragraphs):
                # 添加段落
                para = doc.add_paragraph(translated_para)
                
                # 如果有对应的格式信息，应用格式
                if i < len(styles):
                    style_info = styles[i]
                    
                    # 应用段落样式
                    if 'style_name' in style_info:
                        try:
                            para.style = style_info['style_name']
                        except:
                            pass  # 如果样式不存在，忽略
                    
                    # 应用对齐方式
                    if 'alignment' in style_info:
                        para.alignment = style_info['alignment']
                    
                    # 应用字体信息
                    if 'font' in style_info and para.runs:
                        font_info = style_info['font']
                        run = para.runs[0]
                        
                        if 'name' in font_info and font_info['name']:
                            run.font.name = font_info['name']
                        
                        if 'size' in font_info and font_info['size']:
                            run.font.size = Pt(font_info['size'])
                        
                        if 'bold' in font_info:
                            run.font.bold = font_info['bold']
                        
                        if 'italic' in font_info:
                            run.font.italic = font_info['italic']
                        
                        if 'underline' in font_info:
                            run.font.underline = font_info['underline']
            
            # 保存文档
            doc.save(output_path)
            return True
            
        except Exception as e:
            print(f"创建Word文档失败: {e}")
            return False
    
    def _create_pdf_document(self, original_paragraphs: List[str], translated_paragraphs: List[str], 
                            format_info: Dict[str, Any], output_path: str) -> bool:
        """
        创建格式一致的PDF文档
        
        Args:
            original_paragraphs: 原文段落列表
            translated_paragraphs: 翻译后段落列表
            format_info: 格式信息
            output_path: 输出文件路径
            
        Returns:
            bool: 是否创建成功
        """
        try:
            # 获取页面信息
            pages = format_info.get('pages', [])
            if not pages:
                return self._create_plain_document(translated_paragraphs, output_path, 'pdf')
            
            # 创建PDF文档
            c = canvas.Canvas(output_path)
            
            # 注册中文字体
            try:
                pdfmetrics.registerFont(TTFont('SimSun', 'simsun.ttf'))
                font_name = 'SimSun'
            except:
                font_name = 'Helvetica'  # 如果没有中文字体，使用默认字体
            
            # 处理每一页
            para_index = 0
            for page_info in pages:
                # 设置页面大小
                width = page_info.get('width', 612)  # 默认为letter宽度
                height = page_info.get('height', 792)  # 默认为letter高度
                c.setPageSize((width, height))
                
                # 处理页面上的段落
                page_paragraphs = page_info.get('paragraphs', [])
                for text_info in page_paragraphs:
                    if para_index >= len(translated_paragraphs):
                        break
                    
                    # 获取文本位置和大小
                    x = text_info.get('x0', 50)
                    y = height - text_info.get('y0', 50)  # PDF坐标系与页面坐标系不同
                    size = text_info.get('size', 12)
                    
                    # 设置字体和大小
                    c.setFont(font_name, size)
                    
                    # 绘制文本
                    c.drawString(x, y, translated_paragraphs[para_index])
                    para_index += 1
                
                # 处理图片位置（仅留空白）
                images = format_info.get('images', [])
                for img_info in images:
                    if img_info.get('page', 0) == pages.index(page_info):
                        # 获取图片位置
                        x0 = img_info.get('x0', 0)
                        y0 = height - img_info.get('y0', 0)  # PDF坐标系与页面坐标系不同
                        x1 = img_info.get('x1', 0)
                        y1 = height - img_info.get('y1', 0)
                        
                        # 绘制空白框
                        c.rect(x0, y1, x1-x0, y0-y1)
                
                # 结束当前页
                c.showPage()
            
            # 保存文档
            c.save()
            return True
            
        except Exception as e:
            print(f"创建PDF文档失败: {e}")
            return False
