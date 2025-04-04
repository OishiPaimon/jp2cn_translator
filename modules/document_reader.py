"""
日语到中文翻译程序 - 文档读取模块

此模块负责读取Word和PDF文档，并提取文本内容
"""

import os
from typing import List, Dict, Tuple, Any
import docx
import PyPDF2
import pdfplumber
from pathlib import Path

class DocumentReader:
    """文档读取器，支持Word和PDF格式"""
    
    def __init__(self, file_path: str):
        """
        初始化文档读取器
        
        Args:
            file_path: 文档路径
        """
        self.file_path = file_path
        self.file_extension = os.path.splitext(file_path)[1].lower()
        self.document = None
        self.paragraphs = []
        self.format_info = {}
        
    def read(self) -> Tuple[List[str], Dict[str, Any]]:
        """
        读取文档内容
        
        Returns:
            Tuple[List[str], Dict[str, Any]]: 段落列表和格式信息
        """
        if self.file_extension == '.docx':
            return self._read_docx()
        elif self.file_extension == '.pdf':
            return self._read_pdf()
        else:
            raise ValueError(f"不支持的文件格式: {self.file_extension}")
    
    def _read_docx(self) -> Tuple[List[str], Dict[str, Any]]:
        """
        读取Word文档
        
        Returns:
            Tuple[List[str], Dict[str, Any]]: 段落列表和格式信息
        """
        doc = docx.Document(self.file_path)
        self.document = doc
        
        paragraphs = []
        format_info = {'styles': [], 'images': []}
        
        # 提取段落和格式信息
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                paragraphs.append(para.text)
                # 保存段落样式信息
                style_info = {
                    'index': i,
                    'style_name': para.style.name,
                    'alignment': para.alignment,
                    'font': self._extract_font_info(para),
                }
                format_info['styles'].append(style_info)
        
        # 提取图片位置信息
        image_index = 0
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                format_info['images'].append({
                    'index': image_index,
                    'rel_id': rel.rId,
                    'target': rel.target_ref
                })
                image_index += 1
        
        self.paragraphs = paragraphs
        self.format_info = format_info
        
        return paragraphs, format_info
    
    def _extract_font_info(self, paragraph) -> Dict[str, Any]:
        """
        提取段落的字体信息
        
        Args:
            paragraph: docx段落对象
            
        Returns:
            Dict[str, Any]: 字体信息
        """
        font_info = {}
        
        # 尝试获取第一个运行块的字体信息作为段落字体信息
        if paragraph.runs:
            run = paragraph.runs[0]
            font_info = {
                'name': run.font.name,
                'size': run.font.size,
                'bold': run.font.bold,
                'italic': run.font.italic,
                'underline': run.font.underline,
            }
        
        return font_info
    
    def _read_pdf(self) -> Tuple[List[str], Dict[str, Any]]:
        """
        读取PDF文档
        
        Returns:
            Tuple[List[str], Dict[str, Any]]: 段落列表和格式信息
        """
        # 使用PyPDF2提取文本
        pdf_reader = PyPDF2.PdfReader(self.file_path)
        
        # 使用pdfplumber提取更详细的格式信息
        pdf = pdfplumber.open(self.file_path)
        
        paragraphs = []
        format_info = {
            'pages': [],
            'fonts': [],
            'images': []
        }
        
        # 遍历每一页
        for i, page in enumerate(pdf_reader.pages):
            page_text = page.extract_text()
            if page_text:
                # 按段落分割文本
                page_paragraphs = [p for p in page_text.split('\n\n') if p.strip()]
                paragraphs.extend(page_paragraphs)
                
                # 使用pdfplumber提取页面格式信息
                plumber_page = pdf.pages[i]
                
                # 提取页面信息
                page_info = {
                    'index': i,
                    'width': plumber_page.width,
                    'height': plumber_page.height,
                    'paragraphs': []
                }
                
                # 提取文本块信息
                for j, text_obj in enumerate(plumber_page.extract_words()):
                    text_info = {
                        'index': j,
                        'text': text_obj.get('text', ''),
                        'x0': text_obj.get('x0', 0),
                        'y0': text_obj.get('y0', 0),
                        'x1': text_obj.get('x1', 0),
                        'y1': text_obj.get('y1', 0),
                        'size': text_obj.get('size', 0),
                    }
                    page_info['paragraphs'].append(text_info)
                
                # 提取图片信息
                for j, img in enumerate(plumber_page.images):
                    img_info = {
                        'page': i,
                        'index': j,
                        'x0': img.get('x0', 0),
                        'y0': img.get('y0', 0),
                        'x1': img.get('x1', 0),
                        'y1': img.get('y1', 0),
                    }
                    format_info['images'].append(img_info)
                
                format_info['pages'].append(page_info)
        
        pdf.close()
        
        self.paragraphs = paragraphs
        self.format_info = format_info
        
        return paragraphs, format_info
    
    def get_paragraphs(self) -> List[str]:
        """
        获取文档段落列表
        
        Returns:
            List[str]: 段落列表
        """
        if not self.paragraphs:
            self.read()
        return self.paragraphs
    
    def get_format_info(self) -> Dict[str, Any]:
        """
        获取文档格式信息
        
        Returns:
            Dict[str, Any]: 格式信息
        """
        if not self.format_info:
            self.read()
        return self.format_info
