"""
日语样例文档生成脚本

此脚本用于生成测试用的日语Word和PDF文档
"""

import os
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

# 样例日语文本
SAMPLE_JA_TEXT = """
日本語サンプルテキスト

これは日本語から中国語への翻訳プログラムをテストするためのサンプル文書です。

第1章：はじめに

1.1 背景
自動翻訳技術は近年急速に発展しており、特に大規模言語モデル（LLM）の登場により、翻訳の品質が大幅に向上しています。しかし、専門用語や特定の文脈における正確な翻訳はまだ課題として残っています。

1.2 目的
本プログラムは、日本語の文書を中国語に翻訳する際に、以下の要件を満たすことを目的としています：
- 原文と逐語逐句で対応すること
- 専門用語の一貫した翻訳を保証すること
- 原文の書式を維持すること

第2章：技術的アプローチ

2.1 文書処理
本プログラムはWord文書とPDF文書の両方を処理できます。文書から抽出されたテキストは段落ごとに分割され、翻訳のために準備されます。

2.2 翻訳エンジン
翻訳にはOpenAI APIまたはOllamaモデルを使用します。これらのAIモデルは、コンテキストを理解し、自然な翻訳を生成する能力を持っています。

2.3 用語辞書
プログラムは永続辞書と一時辞書の両方をサポートしています。これにより、特定の用語が文書全体で一貫して翻訳されることを保証します。

第3章：実装の詳細

3.1 モジュール構造
プログラムは以下のモジュールで構成されています：
- 文書読み取りモジュール
- 翻訳モジュール
- 辞書管理モジュール
- 書式保存モジュール

3.2 ワークフロー
1. 文書を読み込み、テキストと書式情報を抽出します
2. 潜在的な専門用語を特定し、一時辞書を作成します
3. ユーザーが辞書を編集できるようにします
4. テキストを段落ごとに翻訳します
5. 翻訳されたテキストを元の書式で出力します

結論

このプログラムは、日本語文書の中国語への正確な翻訳を実現するための総合的なソリューションを提供します。特に長文書や専門的な内容を含む文書の翻訳に適しています。

付録A：技術用語リスト

自動翻訳 - 機械翻訳
大規模言語モデル - 大型言語モデル
専門用語 - 専門用語
文脈 - コンテキスト
書式保存 - フォーマット保持
"""

def create_sample_docx(output_path):
    """
    创建样例Word文档
    
    Args:
        output_path: 输出文件路径
    """
    doc = Document()
    
    # 添加标题
    title = doc.add_heading('日本語サンプル文書', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加段落
    paragraphs = SAMPLE_JA_TEXT.strip().split('\n\n')
    for para in paragraphs:
        if para.startswith('第'):  # 章节标题
            doc.add_heading(para, level=1)
        elif para.startswith('結論'):  # 结论
            p = doc.add_paragraph()
            p.add_run('結論').bold = True
            p.add_run('\n\n' + para[3:])
        elif ':' in para or '：' in para:  # 小节标题
            parts = para.split('：' if '：' in para else ':')
            doc.add_heading(parts[0], level=2)
            if len(parts) > 1:
                doc.add_paragraph(parts[1])
        else:
            doc.add_paragraph(para)
    
    # 保存文档
    doc.save(output_path)
    print(f"已创建样例Word文档: {output_path}")

def create_sample_pdf(output_path):
    """
    创建样例PDF文档
    
    Args:
        output_path: 输出文件路径
    """
    c = canvas.Canvas(output_path, pagesize=letter)
    width, height = letter
    
    # 尝试注册日语字体
    try:
        # 这里使用默认字体，实际使用时可能需要指定日语字体
        font_name = 'Helvetica'
    except:
        font_name = 'Helvetica'
    
    # 添加标题
    c.setFont(font_name, 18)
    c.drawCentredString(width/2, height-50, '日本語サンプル文書')
    
    # 添加段落
    paragraphs = SAMPLE_JA_TEXT.strip().split('\n\n')
    y = height - 100
    
    for para in paragraphs:
        # 章节标题
        if para.startswith('第'):
            c.setFont(font_name, 16)
            c.drawString(50, y, para)
            y -= 30
        # 结论
        elif para.startswith('結論'):
            c.setFont(font_name, 16)
            c.drawString(50, y, '結論')
            y -= 20
            c.setFont(font_name, 12)
            lines = [para[3:][i:i+70] for i in range(0, len(para[3:]), 70)]
            for line in lines:
                c.drawString(50, y, line)
                y -= 15
        # 小节标题
        elif ':' in para or '：' in para:
            parts = para.split('：' if '：' in para else ':')
            c.setFont(font_name, 14)
            c.drawString(50, y, parts[0])
            y -= 20
            if len(parts) > 1:
                c.setFont(font_name, 12)
                lines = [parts[1][i:i+70] for i in range(0, len(parts[1]), 70)]
                for line in lines:
                    c.drawString(50, y, line)
                    y -= 15
        # 普通段落
        else:
            c.setFont(font_name, 12)
            lines = [para[i:i+70] for i in range(0, len(para), 70)]
            for line in lines:
                c.drawString(50, y, line)
                y -= 15
        
        y -= 10
        
        # 如果页面空间不足，创建新页面
        if y < 50:
            c.showPage()
            y = height - 50
    
    # 保存PDF
    c.save()
    print(f"已创建样例PDF文档: {output_path}")

def main():
    """主函数"""
    # 获取样例目录
    script_dir = Path(__file__).parent
    samples_dir = script_dir / "samples"
    
    # 创建样例目录
    os.makedirs(samples_dir, exist_ok=True)
    
    # 创建样例文档
    docx_path = samples_dir / "sample_ja.docx"
    pdf_path = samples_dir / "sample_ja.pdf"
    
    create_sample_docx(str(docx_path))
    create_sample_pdf(str(pdf_path))

if __name__ == "__main__":
    main()
