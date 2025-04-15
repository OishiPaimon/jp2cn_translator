"""
文档生成模块 - 负责将翻译结果生成为Word文档，保留原文档的格式
"""
import os
from typing import List, Optional
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


class DocumentGenerator:
    def __init__(self):
        """初始化文档生成器"""
        pass

    def save_as_text(self, content: str, output_path: str) -> None:
        """
        将内容保存为文本文件
        
        Args:
            content: 要保存的内容
            output_path: 输出文件路径
        """
        try:
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(content)
            print(f"文本文件已保存: {output_path}")
        except Exception as e:
            raise IOError(f"保存文本文件失败: {str(e)}")

    def save_as_word(self, content: str, output_path: str) -> None:
        """
        将内容保存为Word文档，保留所有换行符
        
        Args:
            content: 要保存的内容
            output_path: 输出文件路径
        """
        try:
            # 创建Word文档
            doc = Document()
            
            # 设置默认字体和字号
            style = doc.styles['Normal']
            style.font.name = 'SimSun'  # 宋体
            style.font.size = Pt(12)
            
            # 按行分割内容
            lines = content.split('\n')
            
            # 处理每一行
            current_paragraph = None
            
            for i, line in enumerate(lines):
                if line.strip() == '':
                    # 空行，添加一个空段落
                    doc.add_paragraph()
                    current_paragraph = None
                else:
                    # 非空行
                    if current_paragraph is None:
                        # 创建新段落
                        current_paragraph = doc.add_paragraph(line)
                    else:
                        # 在当前段落中添加一个换行符和文本
                        current_paragraph.add_run('\n' + line)
            
            # 保存文档
            doc.save(output_path)
            print(f"Word文档已保存: {output_path}")
        except Exception as e:
            raise IOError(f"保存Word文档失败: {str(e)}")

    def generate_documents(self, content: str, output_dir: str, base_name: str) -> tuple:
        """
        生成文本文件和Word文档
        
        Args:
            content: 要保存的内容
            output_dir: 输出目录
            base_name: 基本文件名（不含扩展名）
            
        Returns:
            (文本文件路径, Word文档路径)的元组
        """
        # 确保输出目录存在
        os.makedirs(output_dir, exist_ok=True)
        
        # 构建输出文件路径
        txt_path = os.path.join(output_dir, f"{base_name}.txt")
        docx_path = os.path.join(output_dir, f"{base_name}.docx")
        
        # 保存文件
        self.save_as_text(content, txt_path)
        self.save_as_word(content, docx_path)
        
        return txt_path, docx_path


# 测试代码
if __name__ == "__main__":
    # 初始化文档生成器
    generator = DocumentGenerator()
    
    # 测试内容
    test_content = """这是第一段内容。
这是第一段的第二行。

这是第二段内容。
这包含了一些文本。

这是最后一段。"""
    
    # 创建临时输出目录
    os.makedirs("test_output", exist_ok=True)
    
    # 生成文档
    txt_path, docx_path = generator.generate_documents(
        test_content, "test_output", "test_document"
    )
    
    print(f"生成的文件:\n- 文本文件: {txt_path}\n- Word文档: {docx_path}")
